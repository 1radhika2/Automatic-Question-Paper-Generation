###################################################################################################################################

#GUI and MySQL
from tkinter import *
import tkinter.messagebox
import mysql.connector
from PIL import Image,ImageTk
from nltk.stem import *
from tkinter import *
from tkinter import filedialog  #library for path search 
import tkinter.filedialog      #library for path search 
from PIL import Image,ImageTk   #library for image loading
from tkinter.filedialog import askopenfilename      
#Stemming is the process of producing morphological variants of a root/base word.
#A stemming algorithm reduces the words “chocolates”, “chocolatey”, “choco” to the root word “chocolate”.
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
from sklearn.model_selection import train_test_split #split data into train and test sets
from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
from sklearn.svm import SVC# Support Vector Machine
from sklearn.pipeline import Pipeline #pipeline to implement steps in series
from gensim import parsing # To stem data
from joblib import dump, load #save or load model
from sklearn.metrics import accuracy_score
#import pdftotext.py file here
from pdftotext import convert_pdf_to_string
import json

keywords=['cite', 'add', 'acquire', 'analyze', 'abstract', 'appraise',
'define', 'approximate', 'adapt', 'audit', 'animate', 'assess',
'describe', 'articulate', 'allocate', 'blueprint', 'arrange', 'compare',
'draw', 'associate', 'alphabetize', 'breadboard', 'assemble', 'conclude',
'enumerate', 'characterize', 'apply', 'break', 'down', 'budget', 'contrast',
'identify', 'clarify', 'ascertain', 'characterize', 'categorize', 'counsel',
'index', 'classify', 'assign', 'classify', 'code', 'criticize',
'indicate', 'compare', 'attain', 'compare', 'combine', 'critique', 'difference', 'analyse', 'argue',
'label', 'compute', 'avoid', 'confirm', 'compile', 'defend',
'list', 'contrast', 'backup', 'contrast', 'compose', 'determine',
'match', 'convert', 'calculate', 'correlate', 'construct', 'discriminate',
'meet', 'defend', 'capture', 'detect', 'cope', 'estimate', 'differ',
'name', 'describe', 'change', 'diagnose', 'correspond', 'evaluate', 'what', 'how',
'outline', 'detail', 'classify', 'diagram', 'create', 'explain',
'point', 'differentiate', 'complete', 'differentiate', 'cultivate', 'grade',
'quote', 'discuss', 'compute', 'discriminate', 'debug', 'hire',
'read', 'distinguish', 'construct', 'dissect', 'depict', 'interpret','recall', 'elaborate', 'customize', 'distinguish', 'design', 'judge',
'recite', 'estimate', 'demonstrate', 'document', 'develop', 'justify',
'recognize', 'example', 'depreciate', 'ensure', 'devise', 'measure', 'record', 'explain', 'derive', 'examine', 'dictate', 'predict',
'repeat', 'express', 'determine', 'explain', 'enhance', 'prescribe', 'reproduce', 'extend', 'diminish', 'explore', 'explain', 'rank',
'review', 'extrapolate', 'discover', 'figure', 'out', 'facilitate', 'rate',
'select', 'factor', 'draw', 'file', 'format', 'recommend', 'state', 'generalize', 'employ', 'group', 'formulate', 'release', 'study', 'give',
'examine', 'identify', 'generalize', 'select',
'tabulate', 'infer', 'exercise', 'illustrate', 'generate', 'summarize', 'trace', 'interact', 'explore', 'infer', 'handle', 'support',
'write', 'interpolate', 'expose', 'interrupt', 'import', 'test', 'interpret', 'express', 'inventory', 'improve', 'validate', 'observe', 'factor',
'investigate', 'incorporate', 'verify',
'paraphrase', 'figure', 'layout', 'integrate', 'picture', 'graphically', 'graph', 'manage', 'interface',
'predict', 'handle', 'maximize', 'join', 'review', 'illustrate', 'minimize', 'lecture', 'rewrite', 'interconvert', 'optimize', 'model',
'subtract', 'investigate', 'order', 'modify', 'summarize', 'manipulate', 'outline', 'network', 'translate', 'modify', 'point out', 'organize',
'visualize', 'operate', 'prioritize', 'outline',
'personalize', 'proofread', 'overhaul', 'plot', 'query', 'plan', 'practice', 'relate', 'portray', 'predict', 'select', 'prepare',
'separate', 'prescribe', 'price', 'size up', 'produce', 'process', 'subdivide', 'program',
'train', 'rearrange', 'project', 'transform', 'reconstruct',
'protect', 'refer', 'provide', 'relate',  'reorganize',
'round off', 'revise', 'sequence', 'rewrite', 'show', 'specify', 'simulate', 'summarize', 'sketch', 'write', 'solve', 'subscribe'
,'tabulate', 'transcribe', 'translate', 'use']
def uniques(l1):
    l2=[]
    for x in l1:
        if x not in l2:
            l2.append(x)
    return l2
keywords2=[]
keywords2=uniques(keywords)

with open('count_record.json') as f1:
    count_record = json.load(f1)
with open('Knowledge_count.json') as f2:
    Knowledge_count = json.load(f2)
with open('Comprehension_count.json') as f3:
    Comprehension_count = json.load(f3)
with open('Application_count.json') as f4:
    Application_count = json.load(f4)
with open('Analysis_count.json') as f5:
    Analysis_count = json.load(f5)
with open('Synthesis_count.json') as f6:
    Synthesis_count = json.load(f6)
with open('Evaluation_count.json') as f7:
    Evaluation_count = json.load(f7)

def m5():
  R1=Tk()
  R1.geometry('1350x750')
  R1.title('Selection Window')
  
  image=Image.open('l.jpg')
  image=image.resize((1350,695))
  photo_image=ImageTk.PhotoImage(image)
  l=Label(R1, image=photo_image)
  l.place(x=0, y=0)
  
  l=Label(R1, text="QUESTION PREDICTION AND PAPER GENERATION - BLOOM'S TAXONOMY", font=('Times',26,'bold'))
  l.place(x=30, y=60)

  b2=Button(R1, text="SINGLE QUESTION PREDICTION",width=35,height=2,font=('Times',18,'bold'),bg="light salmon", fg="red4", command=Prediction_Single_Question)
  b2.place(x=150, y=200)

  b3=Button(R1, text="MULTIPLE QUESTION PREDICTION",width=35,height=2,font=('Times',18,'bold'),bg="peach puff", fg="brown4", command=Prediction_Multiple_Questions)
  b3.place(x=150, y=300)

  b4=Button(R1, text="ENTER QUESTION BANK",width=35,height=2,font=('Times',18,'bold'),bg="light salmon", fg="red4", command=Enter_Question_Bank)
  b4.place(x=150, y=400)

  b5=Button(R1, text="GENERATE QUESTION PAPER",width=35,height=2,font=('Times',18,'bold'),bg="peach puff", fg="brown4", command=Build_Question_Paper)
  b5.place(x=150, y=500)

  R1.mainloop()


#Data Preprocessing
def SVM_Accuracy():
  print('\n\n---------------------------Accuracy---------------------------------')
  #Accuracy
  #Read the dataset
  df = pd.read_csv("dataset.csv", encoding='latin1')

  #Split the dataset in x and y
  #X, y = df['word'], df['taxonomy']
  X, y = df['Questions'], df['Category']   

  #Split data in train and test sets.
  X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25)

  #Use pipeline to carry out steps in sequence with a single object
  #SVM's rbf kernel gives highest accuracy in this classification problem.
  #Support Vector Machine” (SVM) is a supervised machine learning algorithm which can be used for both classification.
  text_clf = Pipeline([('vect', CountVectorizer()), ('tfidf', TfidfTransformer()), ('clf', SVC(kernel='rbf'))])

  #Train the model
  text_clf.fit(X_train, y_train)
  #print(X_test)
  predicted = text_clf.predict(X_test)

  #Test the model
  #Print the accuracy
  from sklearn.metrics import accuracy_score  #accuracy  
  print('Accuracy:',accuracy_score(y_test, predicted))

#Prediction
def Prediction_Single_Question():
  print('\n\n--------------------------UPLOAD PDF---------------------------------')
  t = Toplevel()
  t.title("UPLOAD PDF")
  t.geometry('750x400')
  t.configure(bg='light blue')
  t.resizable(width = FALSE ,height= FALSE)

  image=Image.open('y.jpg')
  image=image.resize((750,400))
  photo_image=ImageTk.PhotoImage(image)
  l=Label(t, image=photo_image)
  l.place(x=0, y=0)

  def browse():
      path1=tkinter.filedialog.askopenfilename()  #To open file: Dialog that requests selection of an existing file.
      e2.delete(0, END)   #to delete all text in the widget.
      e2.insert(0, path1)   #to insert the text in the widget from the path

  e2 = Entry(t,bd=10,text='',width=70)
  e2.place(x =50 ,y=150)
  #Load joblib file
  text_clf = load('model.joblib')

  def nw():
      path1=e2.get()

  #Convert the pdf to string
      pdf_string = convert_pdf_to_string(path1)
      print(pdf_string)
  #print(type(pdf_string))

  #Prediction
      predicted = text_clf.predict(pdf_string)
      predicted_list = tuple(set(predicted))
  #print(predicted_list)
      weights_dict = {}
      weight1 = 0
      weight2 = 0
      weight3 = 0
      weight4 = 0
      weight5 = 0
      weight6 = 0
      for j in keywords2:
        if j in pdf_string:
          for i in predicted_list:    
            if i==str('Application'):
              weight1 += Application_count[j]/count_record[j]
              weights_dict[i]=weight1
            if i==str('Knowledge'):
              weight2 += Knowledge_count[j]/count_record[j]
              weights_dict[i]=weight2
            if i==str('Analysis'):
              weight3 += Analysis_count[j]/count_record[j]
              weights_dict[i]=weight3
            if i==str('Comprehension'):
              weight4 += Comprehension_count[j]/count_record[j]
              weights_dict[i]=weight4
            if i==str('Synthesis'):
              weight5 += Synthesis_count[j]/count_record[j]
              weights_dict[i]=weight5
            if i==str('Evaluation'):
              weight6 += Evaluation_count[j]/count_record[j]
              weights_dict[i]=weight6
      if (not (weights_dict)):
        for i in predicted_list:
          if i==str('Knowledge'):
            print("Knowledge")
          if i==str('Comprehension'):
            print("Comprehension")
          if i==str('Application'):
            print("Application")
          if i==str('Analysis'):
            print("Analysis")
          if i==str('Synthesis'):
            print("Synthesis")
          if i==str('Evaluation'):
            print("Evaluation")
                            
      else:
        max=0
        for key in weights_dict:
          if weights_dict[key] > max:
            max = weights_dict[key] 
        for key,value in  weights_dict.items():
          if max == value:   
            if key==str('Knowledge'):
              print("Knowledge")
              break
            if key==str('Comprehension'):
              print("Comprehension")
              break
            if key==str('Application'):
              print("Application")
              break
            if key==str('Analysis'):
              print("Analysis")
              break
            if key==str('Synthesis'):
              print("Synthesis")
              break
            if key==str('Evaluation'):
              print("Evaluation")
              break

  b1 = Button(t, text='BROWSE',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=browse)
  b1.place(x =550 ,y=150)
      
  b2 = Button(t, text='OK',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=nw)
  b2.place(x =150 ,y=250)

  label = Label(t, text='UPLOAD PDF', fg='brown4', bg='light salmon', font=('Times',18,'bold'))
  label.place(x=70, y=100)
          
  t.mainloop()
          


def Prediction_Multiple_Questions():
  print('\n\n---------------------------Upload Question Paper---------------------------------')
  import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
  from sklearn.model_selection import train_test_split #split data into train and test sets
  from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
  from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
  from sklearn.svm import SVC# Support Vector Machine
  from sklearn.pipeline import Pipeline #pipeline to implement steps in series
  from gensim import parsing # To stem data
  from joblib import dump, load #save or load model
  from sklearn.metrics import accuracy_score
 

  def question_bank():
      t = Toplevel()
      t.title("UPLOAD QUESTION PAPER")
      t.geometry('750x400')
      t.configure(bg='light blue')
      t.resizable(width = FALSE ,height= FALSE)

      
      image=Image.open('y.jpg')
      image=image.resize((750,400))
      photo_image=ImageTk.PhotoImage(image)
      l=Label(t, image=photo_image)
      l.place(x=0, y=0)
      
         
      def browse():
          path1=tkinter.filedialog.askopenfilename()  #To open file: Dialog that requests selection of an existing file.
          e2.delete(0, END)   #to delete all text in the widget.
          e2.insert(0, path1)   #to insert the text in the widget from the path

      e2 = Entry(t,bd=10,text='',width=70)
      e2.place(x =50 ,y=150)

      def nw():

          #Read the dataset in GUI
          path1 = e2.get()

          import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
          from sklearn.model_selection import train_test_split #split data into train and test sets
          from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
          from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
          from sklearn.svm import SVC# Support Vector Machine
          from sklearn.pipeline import Pipeline #pipeline to implement steps in series
          from gensim import parsing # To stem data
          from joblib import dump, load #save or load model
          from sklearn.metrics import accuracy_score
          from pdftotext import convert_pdf_to_string

          ##############################################################################
          def create_category():    
              import csv
              from csv import DictReader
              #DictReader works by reading in the first line of the CSV
              with open(path1, 'r', errors='ignore') as read_obj:
                  csv_dict_reader = DictReader(read_obj)
                  
                  for row in csv_dict_reader:
                      text=row['Questions']
                      print("\n",text)
                      text1 = text.lower()
                      #Pass the lower case converted variable
                      sentence = str(text1)

                      text_clf = load('model.joblib')

                      from nltk.tokenize import sent_tokenize, word_tokenize
                      token = word_tokenize(text1)
                      
                      token1=[''.join(i)  for i in token]
                      #print('token1',token1)
                      
                      token2=' '.join(token1)
                      #print('token2',token2)
                                                     
                      #Prediction
                      predicted = text_clf.predict(token)
                      predicted_list = tuple(set(predicted))
                      print('predicted_list',predicted_list)

                      weights_dict = {}
                      weight1 = 0
                      weight2 = 0
                      weight3 = 0
                      weight4 = 0
                      weight5 = 0
                      weight6 = 0
                      for j in keywords2:
                        if j in token:
                            for i in predicted_list:    
                                if i==str('Application'):
                                    weight1 += Application_count[j]/count_record[j]
                                    weights_dict[i]=weight1
                                if i==str('Knowledge'):
                                    weight2 += Knowledge_count[j]/count_record[j]
                                    weights_dict[i]=weight2
                                if i==str('Analysis'):
                                    weight3 += Analysis_count[j]/count_record[j]
                                    weights_dict[i]=weight3
                                if i==str('Comprehension'):
                                    weight4 += Comprehension_count[j]/count_record[j]
                                    weights_dict[i]=weight4
                                if i==str('Synthesis'):
                                    weight5 += Synthesis_count[j]/count_record[j]
                                    weights_dict[i]=weight5
                                if i==str('Evaluation'):
                                    weight6 += Evaluation_count[j]/count_record[j]
                                    weights_dict[i]=weight6
                      if (not (weights_dict)):
                        for i in predicted_list:
                            if i==str('Knowledge'):
                                print("Knowledge")
                            if i==str('Comprehension'):
                                print("Comprehension")
                            if i==str('Analysis'):
                                print("Analysis")
                            if i==str('Application'):
                                print("Application")
                            if i==str('Synthesis'):
                                print("Synthesis")
                            if i==str('Evaluation'):
                                print("Evaluation")
                      else:
                        max=0
                        for key in weights_dict:
                            if weights_dict[key] > max:
                                max = weights_dict[key] 
                        for key,value in  weights_dict.items():
                            if max == value:   
                                if key==str('Knowledge'):
                                    print("Knowledge")
                                    break
                                if key==str('Comprehension'):
                                    print("Comprehension")
                                    break
                                if key==str('Analysis'):
                                    print("Analysis")
                                    break
                                if key==str('Application'):
                                    print("Application")
                                    break
                                if key==str('Synthesis'):
                                    print("Synthesis")
                                    break
                                if key==str('Evaluation'):
                                    print("Evaluation")
                                    break

          #print("done")
          create_category()
          ##############################################################################
          
      b1 = Button(t, text='BROWSE',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=browse)
      b1.place(x =550 ,y=150)
      
      b2 = Button(t, text='OK',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=nw)
      b2.place(x =150 ,y=250)

      label = Label(t, text='UPLOAD QUESTION PAPER', fg='brown4', bg='light salmon', font=('Times',18,'bold'))
      label.place(x=70, y=100)
          
      t.mainloop()    

  question_bank()

def Enter_Question_Bank():
  print('\n\n---------------------------Enter_Question_Bank---------------------------------')
  import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
  from sklearn.model_selection import train_test_split #split data into train and test sets
  from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
  from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
  from sklearn.svm import SVC# Support Vector Machine
  from sklearn.pipeline import Pipeline #pipeline to implement steps in series
  from gensim import parsing # To stem data
  from joblib import dump, load #save or load model
  from sklearn.metrics import accuracy_score
 

  def question_bank():
      t = Toplevel()
      t.title("SELECT QUESTION BANK")
      t.geometry('750x400')
      t.configure(bg='light blue')
      t.resizable(width = FALSE ,height= FALSE)

      
      image=Image.open('y.jpg')
      image=image.resize((750,400))
      photo_image=ImageTk.PhotoImage(image)
      l=Label(t, image=photo_image)
      l.place(x=0, y=0)
         
      def browse():
          path1=tkinter.filedialog.askopenfilename()  #To open file: Dialog that requests selection of an existing file.
          e2.delete(0, END)   #to delete all text in the widget.
          e2.insert(0, path1)   #to insert the text in the widget from the path

      e2 = Entry(t,bd=5,text='',width=80)
      e2.place(x =50 ,y=150)

      def nw():

          #Read the dataset in GUI
          path1 = e2.get()

          import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
          from sklearn.model_selection import train_test_split #split data into train and test sets
          from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
          from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
          from sklearn.svm import SVC# Support Vector Machine
          from sklearn.pipeline import Pipeline #pipeline to implement steps in series
          from gensim import parsing # To stem data
          from joblib import dump, load #save or load model
          from sklearn.metrics import accuracy_score
          from pdftotext import convert_pdf_to_string

          ##############################################################################
          def create_category():    
              import csv
              question1 = csv.writer(open("Knowledge.csv", 'w',  errors='ignore', newline=""), delimiter=",")
              question2 = csv.writer(open("Comprehension.csv", 'w',  errors='ignore', newline=""), delimiter=",")
              question3 = csv.writer(open("Analysis.csv", 'w',  errors='ignore', newline=""), delimiter=",")
              question4 = csv.writer(open("Application.csv", 'w',  errors='ignore', newline=""), delimiter=",")
              question5 = csv.writer(open("Synthesis.csv", 'w',  errors='ignore', newline=""), delimiter=",")
              question6 = csv.writer(open("Evaluation.csv", 'w',  errors='ignore', newline=""), delimiter=",")

              from csv import DictReader
              #DictReader works by reading in the first line of the CSV
              with open(path1, 'r', errors='ignore') as read_obj:
                  csv_dict_reader = DictReader(read_obj)
                  
                  for row in csv_dict_reader:
                      text=row['Questions']
                      text1 = text.lower()
                      #Pass the lower case converted variable
                      sentence = str(text1)

                      text_clf = load('model.joblib')

                      from nltk.tokenize import sent_tokenize, word_tokenize
                      token = word_tokenize(text1)
                      
                      token1=[''.join(i)  for i in token]
                      #print('token1',token1)
                      
                      token2=' '.join(token1)
                      #print('token2',token2)
                                                     
                      #Prediction
                      predicted = text_clf.predict(token)
                      predicted_list = tuple(set(predicted))
                      #print('predicted_list',predicted_list)

                      question1 = csv.writer(open("Knowledge.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      question2 = csv.writer(open("Comprehension.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      question3 = csv.writer(open("Analysis.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      question4 = csv.writer(open("Application.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      question5 = csv.writer(open("Synthesis.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      question6 = csv.writer(open("Evaluation.csv", 'a',  errors='ignore', newline=""), delimiter=",")
                      weights_dict = {}
                      weight1 = 0
                      weight2 = 0
                      weight3 = 0
                      weight4 = 0
                      weight5 = 0
                      weight6 = 0
                      for j in keywords2:
                        if j in token:
                            for i in predicted_list:    
                                if i==str('Application'):
                                    weight1 += Application_count[j]/count_record[j]
                                    weights_dict[i]=weight1
                                if i==str('Knowledge'):
                                    weight2 += Knowledge_count[j]/count_record[j]
                                    weights_dict[i]=weight2
                                if i==str('Analysis'):
                                    weight3 += Analysis_count[j]/count_record[j]
                                    weights_dict[i]=weight3
                                if i==str('Comprehension'):
                                    weight4 += Comprehension_count[j]/count_record[j]
                                    weights_dict[i]=weight4
                                if i==str('Synthesis'):
                                    weight5 += Synthesis_count[j]/count_record[j]
                                    weights_dict[i]=weight5
                                if i==str('Evaluation'):
                                    weight6 += Evaluation_count[j]/count_record[j]
                                    weights_dict[i]=weight6
                      if (not (weights_dict)):
                        for i in predicted_list:
                            if i==str('Knowledge'):
                                question1.writerow([text,""])
                            if i==str('Comprehension'):
                                question2.writerow([text,""])
                            if i==str('Analysis'):
                                question3.writerow([text,""])
                            if i==str('Application'):
                                question4.writerow([text,""])
                            if i==str('Synthesis'):
                                question5.writerow([text,""])
                            if i==str('Evaluation'):
                                question6.writerow([text,""])
                      else:
                        max=0
                        for key in weights_dict:
                            if weights_dict[key] > max:
                                max = weights_dict[key] 
                        for key,value in  weights_dict.items():
                            if max == value:   
                                if key==str('Knowledge'):
                                    question1.writerow([text,""])
                                    break
                                if key==str('Comprehension'):
                                    question2.writerow([text,""])
                                    break
                                if key==str('Analysis'):
                                    question3.writerow([text,""])
                                    break
                                if key==str('Application'):
                                    question4.writerow([text,""])
                                    break
                                if key==str('Synthesis'):
                                    question5.writerow([text,""])
                                    break
                                if key==str('Evaluation'):
                                    question6.writerow([text,""])
                                    break

          print("done")
          create_category()
          ##############################################################################
          
      b1 = Button(t, text='BROWSE',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=browse)
      b1.place(x =550 ,y=150)
      
      b2 = Button(t, text='OK',width=10,height=1,font=('Times',15,'bold'),fg="brown4",relief=RAISED,overrelief=RIDGE,bg="peach puff",command=nw)
      b2.place(x =150 ,y=250)

      label = Label(t, text='SELECT QUESTION BANK', fg='brown4', bg='light salmon', font=('Times',18,'bold'))
      label.place(x=70, y=100)
          
      t.mainloop()    

  question_bank()


#Build_Question_Paper
def Build_Question_Paper():
  print('\n\n---------------------------Buidling Question Paper---------------------------------')
  
  import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
  from sklearn.model_selection import train_test_split #split data into train and test sets
  from sklearn.feature_extraction.text import CountVectorizer #convert text comment into a numeric vector
  from sklearn.feature_extraction.text import TfidfTransformer #use TF IDF transformer to change text vector created by count vectorizer
  from sklearn.svm import SVC# Support Vector Machine
  from sklearn.pipeline import Pipeline #pipeline to implement steps in series
  from gensim import parsing # To stem data
  from joblib import dump, load #save or load model
  from sklearn.metrics import accuracy_score
  from pdftotext import convert_pdf_to_string

  def m1():
    def no_of_questions():
      #Application
      a1 = e1.get()
      a10 = float(a1)

      q1 = e7.get()
      q2 = int(q1)

      b1 = a10*q2/100
      b10 = round(b1)
      c1= int(b10)
      print('c1',c1)

      import random
      converted_list = []
      with open("Application.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c1)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      import csv
      question = csv.writer(open("paper.csv", 'w',  errors='ignore', newline=""), delimiter=",")
      question.writerow(['Questions'])

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])


      #Analysis
      a2 = e2.get()
      a20 = float(a2)
      
      q1 = e7.get()
      q2 = int(q1)
      
      b1 = a20*q2/100
      b10 = round(b1)
      c2= int(b10)
      print('c2',c2)
      
      import random
      converted_list = []
      with open("Analysis.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c2)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])


      #Comprehension    
      a3 = e3.get()
      a30 = float(a3)

      q1 = e7.get()
      q2 = int(q1)
      
      b1 = a30*q2/100
      b10 = round(b1)
      c3= int(b10)
      print('c3',c3)

      import random
      converted_list = []
      with open("Comprehension.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c3)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])

      #Evaluation
      a4 = e4.get()
      a40 = float(a4)
      
      q1 = e7.get()
      q2 = int(q1)
      
      b1 = a40*q2/100
      b10 = round(b1)
      c4= int(b10)
      print('c4',c4)

      import random
      converted_list = []
      with open("Evaluation.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c4)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])
      

      #Knowledge
      a5 = e5.get()
      a50 = float(a5)
      
      q1 = e7.get()
      q2 = int(q1)
      
      b1 = a50*q2/100
      b10 = round(b1)
      c5= int(b10)
      print('c5',c5)

      import random
      converted_list = []
      with open("Knowledge.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c5)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])

     
      #Synthesis
      a6 = e6.get()
      a60 = float(a6)
      
      q1 = e7.get()
      q2 = int(q1)
      
      b1 = a60*q2/100
      b10 = round(b1)
      c6= int(b10)
      print('c6',c6)
      
      import random
      converted_list = []
      with open("Synthesis.csv", "rb") as source:
          lines = [line for line in source]
      random_choice = random.sample(lines, c6)
      for element in random_choice:
          e=str(element, 'utf-8')
          stripped_string = e.strip('"')
          print(stripped_string)
          converted_list.append(stripped_string.strip())

      for element in converted_list:

          question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")
          question.writerow([element])
      question = csv.writer(open("paper.csv", 'a',  errors='ignore', newline=""), delimiter=",")  


      ####################################

      #Number of questions input
      q1 = e7.get()
      q2 = int(q1)
      r = q2
      print('No of questions:',r)

      import pandas as pd
      import numpy as np
      import docx
      import csv

      with open('paper.csv','r') as f:
          reader = csv.reader(f,delimiter = ",")
          data = list(reader)
      dataset_len = data
      #print('dataset_length',len(dataset_len))

      #Generate question paper
      i=len(dataset_len)
      df = pd.read_csv('paper.csv',nrows=r)

      # import docx NOT python-docx
      import docx
      # create an instance of a word document
      doc = docx.Document()
      # add a heading of level 0 (largest heading)
      doc.add_heading('                                            Question Paper', 1)
      doc.add_heading('                                          Branch: COMPUTER SCIENCE', 2)
      #doc.add_heading('', 1)
      doc_para = doc.add_paragraph('[Date: 03-05-2021]                                                                                                                         [Marks:100]')
      #italic
      doc_para.add_run('\n\nNote: Answer all the following questions').bold = True
      doc_para.add_run('\n_________________________________________________________________________________________________________').bold = True


      for i in range(0,df.shape[0]):
          for j in range(df.shape[-1]):
              text = str(df.values[i,j])
              i=i+1
              i=str(i)
              t=i+')'+text
              
          doc.add_paragraph(t)
          doc.save('paper.doc')

    ###################################################################################################################
        
    #GUI 
    R2=Toplevel()
    R2.geometry('660x500')
    R2.title('Blooms Taxonomy')

    
    image=Image.open('t.jpg')
    image=image.resize((660,500))
    photo_image=ImageTk.PhotoImage(image)
    l=Label(R2, image=photo_image)
    l.place(x=0, y=0)
    
   

    l=Label(R2, text="ENTER BLOOM'S PERCENTAGE AND NUMBER OF QUESTIONS EXPECTED", font=('Palatino Linotype',12,'bold'), fg="gray4")
    l.place(x=30, y=30)

    l1=Label(R2, text="APPLICATION (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l1.place(x=110, y=80)
    l2=Label(R2, text="ANALYSIS (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l2.place(x=110, y=130)
    l3=Label(R2, text="COMPREHENSION (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l3.place(x=110, y=180)
    l4=Label(R2, text="EVALUATION (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l4.place(x=110, y=230)
    l5=Label(R2, text="KNOWLEDGE (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l5.place(x=110, y=280)
    l6=Label(R2, text="SYNTHESIS (in %)", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l6.place(x=110, y=330)
    
    e1=Entry(R2, font=14)
    e1.place(x=355, y=85)
    e2=Entry(R2, font=14)
    e2.place(x=355, y=135)
    e3=Entry(R2, font=14)
    e3.place(x=355, y=185)
    e4=Entry(R2, font=14)
    e4.place(x=355, y=235)
    e5=Entry(R2, font=14)
    e5.place(x=355, y=285)
    e6=Entry(R2, font=14)
    e6.place(x=355, y=335)
    
    l7=Label(R2, text="NUMBER OF QUESTIONS", font=('Palatino Linotype',14,'bold'), fg="midnight blue")
    l7.place(x=80, y=390)
    e7=Entry(R2, font=16)
    e7.place(x=380, y=395)

    b1=Button(R2, text="OK",width=8,height=1, font=('Palatino Linotype',14,'bold'), bg="lightblue", fg="black", command=no_of_questions)
    b1.place(x=400, y=430)
      
    R2.mainloop()

  m1()
  
m5()

###################################################################################################################################
